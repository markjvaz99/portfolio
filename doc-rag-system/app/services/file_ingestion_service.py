import os
import tempfile
from typing import List

from fastapi import UploadFile

from llama_index.core import VectorStoreIndex, Document
from llama_index.vector_stores.qdrant import QdrantVectorStore
from llama_index.embeddings.ollama import OllamaEmbedding
from llama_index.core import StorageContext

from qdrant_client import QdrantClient

import fitz  # PyMuPDF
from docx import Document as DocxDocument  # python-docx

from app.config import *
from app.utils.hashing import file_hash
from app.services.file_tracker import init_db, is_processed, mark_processed


# =========================================================
# 🔹 COMMON UTILITIES
# =========================================================

def clean_text(text: str) -> str:
    text = text.replace("\n", " ")
    text = text.replace("  ", " ")
    return text.strip()


def chunk_text(text: str, chunk_size=500, overlap=75):
    chunks = []
    start = 0

    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start += (chunk_size - overlap)

    return chunks


# =========================================================
# 🔹 PDF LOADER (PyMuPDF)
# =========================================================

def load_pdf(file_path: str, filename: str) -> List[Document]:
    documents = []

    doc = fitz.open(file_path)

    for page_num, page in enumerate(doc, start=1):
        text = page.get_text("text")
        text = clean_text(text)

        if not text:
            continue

        chunks = chunk_text(text)

        for i, chunk in enumerate(chunks):
            documents.append(
                Document(
                    text=chunk,
                    metadata={
                        "filename": filename,
                        "type": "pdf",
                        "page": page_num,
                        "chunk_id": i
                    }
                )
            )

    return documents


# =========================================================
# 🔹 DOCX LOADER (python-docx)
# =========================================================

def load_docx(file_path: str, filename: str) -> List[Document]:
    documents = []

    doc = DocxDocument(file_path)

    current_section = "unknown"
    full_text = []

    for para in doc.paragraphs:
        text = para.text.strip()

        if not text:
            continue

        # Detect headings
        if para.style.name.startswith("Heading"):
            current_section = text

        full_text.append((current_section, text))

    # Merge and chunk
    combined_text = "\n".join([t[1] for t in full_text])
    combined_text = clean_text(combined_text)

    chunks = chunk_text(combined_text)

    for i, chunk in enumerate(chunks):
        documents.append(
            Document(
                text=chunk,
                metadata={
                    "filename": filename,
                    "type": "docx",
                    "section": current_section,
                    "chunk_id": i
                }
            )
        )

    return documents


# =========================================================
# 🔹 TXT LOADER
# =========================================================

def load_txt(content: bytes, filename: str) -> List[Document]:
    documents = []

    text = content.decode("utf-8", errors="ignore")
    text = clean_text(text)

    chunks = chunk_text(text)

    for i, chunk in enumerate(chunks):
        documents.append(
            Document(
                text=chunk,
                metadata={
                    "filename": filename,
                    "type": "txt",
                    "chunk_id": i
                }
            )
        )

    return documents


# =========================================================
# 🔹 FILE ROUTER (Loader Dispatcher)
# =========================================================

def load_file(tmp_path: str, content: bytes, filename: str) -> List[Document]:
    filename_lower = filename.lower()

    if filename_lower.endswith(".pdf"):
        return load_pdf(tmp_path, filename)

    elif filename_lower.endswith(".docx"):
        return load_docx(tmp_path, filename)

    elif filename_lower.endswith(".txt"):
        return load_txt(content, filename)

    else:
        return []


# =========================================================
# 🔹 MAIN INGEST FUNCTION
# =========================================================

async def ingest_files(files: List[UploadFile]) -> dict:
    init_db()

    client = QdrantClient(url=QDRANT_URL)

    vector_store = QdrantVectorStore(
        client=client,
        collection_name=COLLECTION_NAME
    )

    storage_context = StorageContext.from_defaults(
        vector_store=vector_store
    )

    embed_model = OllamaEmbedding(
        model_name=EMBED_MODEL,
        base_url=OLLAMA_BASE_URL
    )

    all_documents = []
    processed_files = []
    skipped_files = []

    for file in files:
        # 🔹 Save temp file
        with tempfile.NamedTemporaryFile(delete=False) as tmp:
            content = await file.read()
            tmp.write(content)
            tmp_path = tmp.name

        f_hash = file_hash(tmp_path)

        if is_processed(file.filename, f_hash):
            skipped_files.append(file.filename)
            os.remove(tmp_path)
            continue

        try:
            documents = load_file(tmp_path, content, file.filename)

            if not documents:
                skipped_files.append(file.filename)
                os.remove(tmp_path)
                continue

            all_documents.extend(documents)

        except Exception:
            skipped_files.append(file.filename)
            os.remove(tmp_path)
            continue

        mark_processed(file.filename, f_hash)
        processed_files.append(file.filename)

        os.remove(tmp_path)

    # 🔹 Indexing (shared pipeline)
    if all_documents:
        VectorStoreIndex.from_documents(
            all_documents,
            storage_context=storage_context,
            embed_model=embed_model
        )

    return {
        "processed": processed_files,
        "skipped": skipped_files,
        "total_ingested": len(processed_files)
    }
